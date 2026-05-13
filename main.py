from __future__ import annotations

from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from flask_socketio import SocketIO, emit
from openai import OpenAI
from docx import Document
from collections import deque
from datetime import datetime
import logging
import os
import io
import json
import re
import subprocess
import time
import threading
import config
from services.asr_service import (
    ASRAdapter,
    ChunkDecision,
    RealtimeChunkPolicy,
    SegmentRewritePolicy,
    add_wav_header,
    build_realtime_chunk_policy,
    is_effective_text_update,
    load_realtime_chunk_policy_overrides,
    decide_chunk_processing,
    decide_chunk_processing_simple,
    decide_segment_rewrite,
    decide_stop_flush,
    decide_stop_flush_simple,
    extract_audio_features,
    format_asr_display_text,
    refine_asr_result_text,
    retain_realtime_buffer,
    should_filter_asr_result,
)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("VoiceSystem")

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'voice-system-secret')

socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

client = OpenAI(api_key=config.API_KEY, base_url=config.BASE_URL)

asr_adapter = ASRAdapter(
    asr_mode=config.ASR_MODE,
    asr_base_url=config.ASR_BASE_URL,
    asr_model=config.ASR_MODEL,
    asr_api_key=config.ASR_API_KEY,
)

TEMP_DIR = config.TEMP_DIR
EXPORT_DIR = config.EXPORT_DIR
for d in [TEMP_DIR, EXPORT_DIR]:
    os.makedirs(d, exist_ok=True)

REALTIME_DEBUG_TRACE = deque(maxlen=200)


def env_flag(name: str, default: bool = False) -> bool:
    raw_value = os.getenv(name)
    if raw_value is None:
        return default
    return str(raw_value).strip().lower() in {"1", "true", "yes", "on"}


def env_float(name: str, default: float) -> float:
    raw_value = os.getenv(name)
    if raw_value is None:
        return default
    try:
        return float(str(raw_value).strip())
    except (TypeError, ValueError):
        logger.warning("Invalid float env %s=%r, falling back to %.2f", name, raw_value, default)
        return default


def get_app_js_version() -> int:
    app_js_path = os.path.join(app.static_folder or "static", "js", "app.js")
    try:
        return int(os.path.getmtime(app_js_path))
    except OSError:
        return int(time.time())


def append_realtime_debug_trace(event_type: str, payload: dict) -> None:
    trace_payload = {
        "ts": datetime.now().isoformat(timespec="seconds"),
        "event_type": event_type,
        "result_id": payload.get("result_id"),
        "result_type": payload.get("result_type"),
        "segment_id": payload.get("segment_id"),
        "replace_target_id": payload.get("replace_target_id"),
        "processing_reason": payload.get("processing_reason"),
        "text": payload.get("text"),
    }
    REALTIME_DEBUG_TRACE.append(trace_payload)


def get_git_head_short() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=os.path.dirname(os.path.abspath(__file__)),
            capture_output=True,
            text=True,
            check=False,
            timeout=2,
        )
    except Exception:
        return ""
    if result.returncode != 0:
        return ""
    return result.stdout.strip()

def remove_markdown_formatting(text: str) -> str:
    if not text:
        return text
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'^(#{1,6})\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^[\*\-\+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    return text
def build_transcript_text(transcript, max_chars=10000):
    full_text_list = []
    current_chars = 0
    for t in reversed(transcript):
        entry = f"{t['speaker']}: {t['content']}"
        if current_chars + len(entry) > max_chars:
            full_text_list.append("...[此处省略部分过长内容]...")
            break
        full_text_list.append(entry)
        current_chars += len(entry)
    return "\n".join(reversed(full_text_list))


def parse_doc_response(raw_content, doc_type):
    doc_title = "未命名文档"
    doc_content = raw_content

    if "【标题】" in raw_content:
        parts = raw_content.split("【标题】", 1)
        if len(parts) > 1:
            title_part = parts[1].split("【正文】", 1) if "【正文】" in parts[1] else parts[1].split("\n", 1)
            doc_title = title_part[0].strip()
            if "【正文】" in parts[1]:
                doc_content = parts[1].split("【正文】", 1)[1].strip()
            elif len(title_part) > 1:
                doc_content = title_part[1].strip() if len(title_part) > 1 else ""
            else:
                doc_content = ""
    else:
        doc_type_names = {"meeting": "会议纪要", "report": "出差报告", "publicity": "宣传报道"}
        doc_title = f"{doc_type_names.get(doc_type, '文档')} - {datetime.now().strftime('%Y%m%d %H:%M')}"

    return doc_title, doc_content


DOC_PROMPTS = {
    "meeting": {
        "system": "你是一个专业的会议速记员，负责根据转写内容生成会议纪要。请保持专业、简洁、准确。",
        "user": """请根据以下会议转写内容，生成一份详细且结构化的会议纪要。

【重要约束】
1. 基于转写中的事实进行总结和归纳，**禁止编造转写中完全没有提及的事件或结论**
2. 可以对已有信息进行合理扩充：补充背景说明、用更丰富的语言重述观点、梳理逻辑关系
3. 如果某类信息在转写中确实未提及，可合理推断（如根据讨论内容推断会议目的）或省略

要求：
1. 【文档标题】请先为这份会议纪要生成一个简洁准确的标题（10-20字）
2. 会议主题：简要概括会议目的和背景
3. 核心议题：列出会议讨论的主要点，并对每个议题进行适当展开说明
4. 讨论要点：详细记录各方观点和意见，用完整的句子表达，不要只是简单罗列
5. 待办事项：明确后续需要跟进的任务及责任人（如果有）
6. 关键结论：总结会议达成的共识、决策或主要观点

请按照以下格式输出，内容要充实饱满：
【标题】XXX
【正文】
...会议纪要内容...

转写内容如下：
{full_text}"""
    },
    "report": {
        "system": "你是一个专业的商务文案撰写专家，负责根据转写内容生成出差报告。请保持专业、简洁、全面。",
        "user": """请根据以下出差相关的转写内容，生成一份详细的出差报告。

【重要约束】
1. 基于转写中的事实进行总结和归纳，**禁止编造转写中完全没有提及的事件或成果**
2. 可以对已有信息进行合理扩充：补充背景说明、用更专业的语言重述、梳理工作脉络
3. 如果某类信息在转写中确实未提及，可合理推断或省略

要求：
1. 【文档标题】请先为这份出差报告生成一个简洁准确的标题（10-20字）
2. 出差概况：包括出差背景、目的和意义
3. 主要活动：按时间顺序记录出差期间的主要行程和活动，对重要活动进行详细描述
4. 工作成果：详细总结出差期间达成的工作目标、收获和进展，展开说明成果的价值
5. 问题与建议：记录出差过程中遇到的问题及改进建议（如果有）
6. 后续跟进：列出需要后续跟进的事项和计划（如果有）

请按照以下格式输出，内容要充实饱满：
【标题】XXX
【正文】
...出差报告内容...

转写内容如下：
{full_text}"""
    },
    "publicity": {
        "system": "你是一个专业的宣传文案专家，负责根据转写内容生成宣传报道。请保持积极正面、吸引人、易于传播。",
        "user": """请根据以下转写内容，生成一份面向公众的宣传报道或新闻稿。

【重要约束】
1. 基于转写中的事实进行宣传创作，**禁止编造转写中完全没有提及的数据或事件**
2. 可以对已有信息进行合理扩充：用更生动的语言描述、提炼亮点、添加积极正面的评价和展望
3. 在保持事实真实的前提下，最大化内容的吸引力和感染力

要求：
1. 【文档标题】请先为这份宣传报道生成一个吸引人的标题（10-20字）
2. 导语：用简洁有力的语言概括核心亮点
3. 主体内容：详细阐述转写中提到的亮点和成果，展开说明其意义和价值，用生动的语言描述
4. 语言风格：生动活泼、富有感染力，使用正面积极的词汇和表达
5. 结尾：总结升华，传递正能量，适当展望未来

请按照以下格式输出，内容要充实饱满：
【标题】XXX
【正文】
...宣传报道内容...

转写内容如下：
{full_text}"""
    }
}


@app.route("/")
def index():
    return render_template("index.html", app_js_version=get_app_js_version())


@app.route("/logo.jpeg")
def logo():
    return send_file("logo.jpeg")


@app.route("/api/v1/debug/status")
def get_status():
    return jsonify({
        "status": "ok",
        "mode": "intranet" if config.USE_INTRANET else "online",
        "asr_model": config.ASR_MODEL,
        "llm_model": config.LLM_MODEL,
        "base_url": config.BASE_URL,
        "app_js_version": get_app_js_version(),
        "git_head": get_git_head_short(),
        "realtime_trace_count": len(REALTIME_DEBUG_TRACE),
    })


@app.route("/api/v1/debug/realtime_trace")
def get_realtime_trace():
    return jsonify({
        "count": len(REALTIME_DEBUG_TRACE),
        "events": list(REALTIME_DEBUG_TRACE),
    })


DEFAULT_ENABLE_SIMPLIFIED_REALTIME_PIPELINE = not config.USE_INTRANET
ENABLE_SIMPLIFIED_REALTIME_PIPELINE = env_flag(
    "ENABLE_SIMPLIFIED_REALTIME_PIPELINE",
    DEFAULT_ENABLE_SIMPLIFIED_REALTIME_PIPELINE,
)

DEFAULT_REALTIME_CHUNK_POLICY = build_realtime_chunk_policy(
    simplified=ENABLE_SIMPLIFIED_REALTIME_PIPELINE,
)
REALTIME_CHUNK_POLICY, REALTIME_CHUNK_POLICY_OVERRIDES = load_realtime_chunk_policy_overrides(
    DEFAULT_REALTIME_CHUNK_POLICY,
    os.environ,
    mode_prefix="INTRANET" if config.USE_INTRANET else "ONLINE",
)
SEGMENT_REWRITE_POLICY = SegmentRewritePolicy()
DECIDE_REALTIME_CHUNK = decide_chunk_processing_simple if ENABLE_SIMPLIFIED_REALTIME_PIPELINE else decide_chunk_processing
DECIDE_STOP_FLUSH = decide_stop_flush_simple if ENABLE_SIMPLIFIED_REALTIME_PIPELINE else decide_stop_flush
MEDIUM_REWRITE_SECONDS = env_float("MEDIUM_REWRITE_SECONDS", 6.0)
HIGH_REWRITE_SECONDS = env_float("HIGH_REWRITE_SECONDS", 30.0)
IDLE_SEGMENT_SPLIT_SECONDS = env_float("IDLE_SEGMENT_SPLIT_SECONDS", 2.0)
IDLE_HIGH_REWRITE_SECONDS = env_float("IDLE_HIGH_REWRITE_SECONDS", 10.0)


class SessionManager:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._sessions = {}
            cls._instance._recently_stopped = {}
            cls._instance._lock = threading.Lock()
        return cls._instance

    def get_or_create(self, sid):
        with self._lock:
            if sid not in self._sessions:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"stream_recording_{timestamp}.pcm"
                file_path = os.path.join(TEMP_DIR, filename)
                self._sessions[sid] = {
                    'buffer': bytearray(),
                    'last_process_time': time.time(),
                    'last_audio_time': time.time(),
                    'last_speech_time': 0.0,
                    'last_idle_rewrite_audio_time': 0.0,
                    'processing': False,
                    'stop_requested': False,
                    'drain_lock': threading.Lock(),
                    'file_path': file_path,
                    'file_handle': open(file_path, "wb"),
                    'session_tag': timestamp,
                    'chunk_seq': 0,
                    'result_seq': 0,
                    'segment_seq': 0,
                    'active_segment': None,
                }
            return self._sessions[sid]

    def get(self, sid):
        with self._lock:
            return self._sessions.get(sid)

    def items_snapshot(self):
        with self._lock:
            return list(self._sessions.items())

    def mark_recently_stopped(self, sid, cooldown_seconds=3.0):
        with self._lock:
            self._recently_stopped[sid] = time.time() + cooldown_seconds

    def clear_recently_stopped(self, sid):
        with self._lock:
            self._recently_stopped.pop(sid, None)

    def should_ignore_audio(self, sid):
        with self._lock:
            expires_at = self._recently_stopped.get(sid)
            if expires_at is None:
                return False
            if expires_at < time.time():
                self._recently_stopped.pop(sid, None)
                return False
            return True

    def remove(self, sid):
        with self._lock:
            session = self._sessions.pop(sid, None)
            if session:
                fh = session.get('file_handle')
                if fh:
                    fh.close()
                return session.get('file_path', '')
            return ''


session_mgr = SessionManager()


def transcribe_realtime_chunk(audio_data: bytes) -> str:
    wav_data = add_wav_header(audio_data)
    return asr_adapter.transcribe_audio_bytes(wav_data, filename="audio.wav", mime_type="audio/wav")


def contains_meaningful_realtime_activity(audio_data: bytes) -> bool:
    if not audio_data:
        return False

    features = extract_audio_features(
        audio_data,
        sample_rate=REALTIME_CHUNK_POLICY.sample_rate,
        bytes_per_sample=REALTIME_CHUNK_POLICY.bytes_per_sample,
        frame_size=REALTIME_CHUNK_POLICY.frame_size,
        silence_threshold=REALTIME_CHUNK_POLICY.silence_threshold,
        peak_threshold=REALTIME_CHUNK_POLICY.peak_threshold,
        speech_rms_threshold=REALTIME_CHUNK_POLICY.speech_rms_threshold,
        speech_peak_threshold=REALTIME_CHUNK_POLICY.speech_peak_threshold,
    )
    if features.duration_seconds <= 0:
        return False

    min_packet_rms = max(
        REALTIME_CHUNK_POLICY.silence_threshold * 1.8,
        REALTIME_CHUNK_POLICY.active_speech_rms_threshold * 0.65,
    )
    min_packet_peak = max(
        REALTIME_CHUNK_POLICY.peak_threshold + 40,
        int(REALTIME_CHUNK_POLICY.active_speech_peak_threshold * 0.65),
    )
    min_active_run_seconds = max(0.03, REALTIME_CHUNK_POLICY.min_active_run_seconds * 0.2)
    min_voiced_run_seconds = max(0.02, REALTIME_CHUNK_POLICY.min_voiced_run_seconds * 0.35)

    has_packet_energy = (
        features.rms >= min_packet_rms
        and features.peak >= min_packet_peak
    )
    has_packet_presence = (
        features.active_ratio >= 0.15
        or features.max_active_run_seconds >= min_active_run_seconds
        or features.max_voiced_run_seconds >= min_voiced_run_seconds
    )
    return has_packet_energy and has_packet_presence


def get_or_create_active_segment(session):
    active_segment = session.get('active_segment')
    if active_segment is None:
        session['segment_seq'] = session.get('segment_seq', 0) + 1
        active_segment = {
            "segment_id": f"{session.get('session_tag', 'session')}_segment_{session['segment_seq']}",
            "audio_buffer": bytearray(),
            "chunk_count": 0,
            "duration_seconds": 0.0,
            "stage_chunk_count": 0,
            "stage_duration_seconds": 0.0,
            "stable_text": "",
            "stage_display_text": "",
            "last_medium_text": "",
            "last_result_id": None,
            "latest_display_text": "",
            "latest_result_type": None,
            "next_medium_rewrite_seconds": MEDIUM_REWRITE_SECONDS,
        }
        session['active_segment'] = active_segment
    return active_segment


def build_realtime_result_payload(
    session,
    text_result: str,
    chunk_decision,
    *,
    segment_id: str | None = None,
    replace_target_id: str | None = None,
    result_type: str = "final_chunk",
    processing_reason: str | None = None,
    stable_text: str = "",
    medium_text: str = "",
    partial_text: str = "",
):
    session['result_seq'] += 1
    chunk_seq = session.get('chunk_seq', 0)
    result_seq = session['result_seq']
    session_tag = session.get('session_tag', 'session')
    return {
        "speaker_id": "Speaker_1",
        "text": text_result,
        "is_final": True,
        "result_id": f"{session_tag}_result_{result_seq}",
        "segment_id": segment_id or f"{session_tag}_segment_{chunk_seq}",
        "replace_target_id": replace_target_id,
        "result_type": result_type,
        "processing_reason": processing_reason or chunk_decision.reason,
        "chunk_duration_seconds": round(chunk_decision.audio_duration_seconds, 3),
        "stable_text": stable_text,
        "medium_text": medium_text,
        "partial_text": partial_text,
    }


def cleanup_realtime_session(sid):
    file_path = session_mgr.remove(sid)
    if file_path:
        logger.info(f"录音结束，音频已保存至: {file_path}")


def emit_realtime_result(sid, payload: dict) -> None:
    socketio.emit("asr_result", payload, to=sid)


def merge_transcript_fragments(previous_text: str, next_text: str) -> str:
    previous = (previous_text or "").strip()
    next_value = (next_text or "").strip()
    if not previous:
        return next_value
    if not next_value:
        return previous

    previous_char = previous[-1]
    next_char = next_value[0]
    if previous_char.isalnum() and next_char.isalnum():
        return f"{previous} {next_value}"
    return f"{previous}{next_value}"


def build_segment_display_text(active_segment, stage_text: str | None = None) -> str:
    stable_text = active_segment.get("stable_text", "")
    suffix_text = active_segment.get("stage_display_text", "") if stage_text is None else stage_text
    return merge_transcript_fragments(stable_text, suffix_text)


def has_displayable_segment_content(active_segment) -> bool:
    if not active_segment:
        return False
    return bool(
        active_segment.get("stage_display_text")
        or active_segment.get("stable_text")
        or active_segment.get("latest_display_text")
    )


def build_display_layers(active_segment, result_type: str, *, stage_text: str = "", stable_override: str | None = None) -> tuple[str, str, str]:
    stable_text = active_segment.get("stable_text", "") if stable_override is None else stable_override
    if result_type == "high_rewrite":
        return stable_text, "", ""
    if result_type == "medium_rewrite":
        return stable_text, stage_text, ""

    medium_text = active_segment.get("last_medium_text", "")
    if stage_text and medium_text and stage_text.startswith(medium_text):
        return stable_text, medium_text, stage_text[len(medium_text):]
    return stable_text, "", stage_text


def should_emit_realtime_update(active_segment, display_text: str, result_type: str) -> bool:
    latest_display_text = active_segment.get("latest_display_text", "")
    latest_result_type = active_segment.get("latest_result_type")
    return latest_result_type != result_type or is_effective_text_update(latest_display_text, display_text)


def emit_realtime_display_update(
    sid,
    session,
    active_segment,
    chunk_decision,
    *,
    display_text: str,
    result_type: str,
    processing_reason: str,
    stable_text: str = "",
    medium_text: str = "",
    partial_text: str = "",
) -> bool:
    if not display_text:
        return False

    if not should_emit_realtime_update(active_segment, display_text, result_type):
        logger.info(
            "Skipping noop realtime update(segment=%s, type=%s, reason=%s)",
            active_segment['segment_id'],
            result_type,
            processing_reason,
        )
        return False

    realtime_payload = build_realtime_result_payload(
        session,
        display_text,
        chunk_decision,
        segment_id=active_segment['segment_id'],
        replace_target_id=active_segment.get('last_result_id'),
        result_type=result_type,
        processing_reason=processing_reason,
        stable_text=stable_text,
        medium_text=medium_text,
        partial_text=partial_text,
    )
    append_realtime_debug_trace("emit_asr_result", realtime_payload)
    emit_realtime_result(sid, realtime_payload)
    active_segment['last_result_id'] = realtime_payload["result_id"]
    active_segment['latest_display_text'] = display_text
    active_segment['latest_result_type'] = result_type
    return True


def reset_active_stage(active_segment) -> None:
    active_segment['audio_buffer'] = bytearray()
    active_segment['stage_chunk_count'] = 0
    active_segment['stage_duration_seconds'] = 0.0
    active_segment['stage_display_text'] = ""
    active_segment['last_medium_text'] = ""
    active_segment['next_medium_rewrite_seconds'] = MEDIUM_REWRITE_SECONDS


def emit_stage_rewrite(
    session,
    sid,
    active_segment,
    chunk_decision,
    *,
    result_type: str,
    processing_reason: str,
    finalize_stage: bool = False,
    ensure_sentence_end: bool = False,
) -> bool:
    stage_audio = bytes(active_segment.get('audio_buffer', b""))
    if not stage_audio and not finalize_stage:
        return False

    stage_text = active_segment.get('stage_display_text', '')
    raw_text_result = ""
    refined_text_result = ""

    if stage_audio:
        raw_text_result = transcribe_realtime_chunk(stage_audio)
        refined_text_result = refine_asr_result_text(raw_text_result)

    if refined_text_result and not should_filter_asr_result(refined_text_result):
        stage_text = format_asr_display_text(
            raw_text_result,
            ensure_sentence_end=ensure_sentence_end,
        ) or refined_text_result
    elif raw_text_result:
        logger.info(
            "ASR rewrite filtered(segment=%s, type=%s, reason=%s, raw=%s, refined=%s)",
            active_segment['segment_id'],
            result_type,
            processing_reason,
            raw_text_result[:50],
            refined_text_result[:50] if refined_text_result else "",
        )

    if finalize_stage:
        committed_text = build_segment_display_text(active_segment, stage_text)
        stable_text, medium_text, partial_text = build_display_layers(
            active_segment,
            result_type,
            stage_text="",
            stable_override=committed_text,
        )
        emitted = emit_realtime_display_update(
            sid,
            session,
            active_segment,
            chunk_decision,
            display_text=committed_text,
            result_type=result_type,
            processing_reason=processing_reason,
            stable_text=stable_text,
            medium_text=medium_text,
            partial_text=partial_text,
        )
        active_segment['stable_text'] = committed_text
        reset_active_stage(active_segment)
        logger.info(
            "ASR high rewrite committed(segment=%s, reason=%s, total_chunks=%s, total_duration=%.2fs, text=%s)",
            active_segment['segment_id'],
            processing_reason,
            active_segment['chunk_count'],
            active_segment['duration_seconds'],
            committed_text[:80],
        )
        return emitted

    if not stage_text:
        return False

    combined_display_text = build_segment_display_text(active_segment, stage_text)
    stable_text, medium_text, partial_text = build_display_layers(
        active_segment,
        result_type,
        stage_text=stage_text,
    )
    emitted = emit_realtime_display_update(
        sid,
        session,
        active_segment,
        chunk_decision,
        display_text=combined_display_text,
        result_type=result_type,
        processing_reason=processing_reason,
        stable_text=stable_text,
        medium_text=medium_text,
        partial_text=partial_text,
    )
    active_segment['stage_display_text'] = stage_text
    if result_type == "medium_rewrite":
        active_segment['last_medium_text'] = stage_text
    logger.info(
        "ASR rewrite emitted(segment=%s, type=%s, reason=%s, stage_duration=%.2fs, text=%s)",
        active_segment['segment_id'],
        result_type,
        processing_reason,
        active_segment['stage_duration_seconds'],
        combined_display_text[:80],
    )
    return emitted


def emit_tiered_rewrite_if_needed(session, sid, active_segment, chunk_decision, *, finalize_segment=False):
    if not active_segment:
        return False

    rewrite_emitted = False

    while active_segment.get('stage_duration_seconds', 0.0) >= HIGH_REWRITE_SECONDS:
        rewrite_emitted = emit_stage_rewrite(
            session,
            sid,
            active_segment,
            chunk_decision,
            result_type="high_rewrite",
            processing_reason="high_rewrite_window_30s",
            finalize_stage=True,
            ensure_sentence_end=False,
        ) or rewrite_emitted

    while (
        active_segment.get('stage_duration_seconds', 0.0) >= active_segment.get('next_medium_rewrite_seconds', MEDIUM_REWRITE_SECONDS)
        and active_segment.get('stage_duration_seconds', 0.0) < HIGH_REWRITE_SECONDS
    ):
        threshold_seconds = active_segment.get('next_medium_rewrite_seconds', MEDIUM_REWRITE_SECONDS)
        rewrite_emitted = emit_stage_rewrite(
            session,
            sid,
            active_segment,
            chunk_decision,
            result_type="medium_rewrite",
            processing_reason=f"medium_rewrite_window_{int(threshold_seconds)}s",
            finalize_stage=False,
            ensure_sentence_end=False,
        ) or rewrite_emitted
        active_segment['next_medium_rewrite_seconds'] = threshold_seconds + MEDIUM_REWRITE_SECONDS

    if finalize_segment and (
        active_segment.get('stage_duration_seconds', 0.0) > 0
        or active_segment.get('stable_text')
    ):
        rewrite_emitted = emit_stage_rewrite(
            session,
            sid,
            active_segment,
            chunk_decision,
            result_type="high_rewrite",
            processing_reason="stop_recording_finalize_segment",
            finalize_stage=True,
            ensure_sentence_end=True,
        ) or rewrite_emitted

    return rewrite_emitted


def process_realtime_audio_chunk(session, sid, audio_data: bytes, chunk_decision, *, finalize_after_processing=False):
    session['last_process_time'] = time.time()
    session['processing'] = True
    session['chunk_seq'] += 1
    active_segment = get_or_create_active_segment(session)
    active_segment['audio_buffer'].extend(audio_data)
    active_segment['chunk_count'] += 1
    active_segment['duration_seconds'] += chunk_decision.audio_duration_seconds
    active_segment['stage_chunk_count'] += 1
    active_segment['stage_duration_seconds'] += chunk_decision.audio_duration_seconds

    logger.info(
        "Processing realtime chunk sid=%s chunk=%s reason=%s speech_gate=%s tail_gate=%s duration=%.2fs bytes=%s rms=%.4f peak=%s active=%.2f voiced=%.2f density=%.2f active_s=%.2fs voiced_s=%.2fs active_run_s=%.2fs voiced_run_s=%.2fs silence=%.2f",
        sid,
        session['chunk_seq'],
        chunk_decision.reason,
        chunk_decision.speech_gate_reason or "-",
        chunk_decision.tail_gate_reason or "-",
        chunk_decision.audio_duration_seconds,
        len(audio_data),
        chunk_decision.audio_features.rms if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.peak if chunk_decision.audio_features else 0,
        chunk_decision.audio_features.active_ratio if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.voiced_ratio if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.voiced_density if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.active_seconds if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.voiced_seconds if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.max_active_run_seconds if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.max_voiced_run_seconds if chunk_decision.audio_features else 0.0,
        chunk_decision.audio_features.silence_ratio if chunk_decision.audio_features else 0.0,
    )

    try:
        text_result = transcribe_realtime_chunk(audio_data)
        refined_text_result = refine_asr_result_text(text_result)
        is_filtered_result = not refined_text_result or should_filter_asr_result(refined_text_result)

        if refined_text_result and not is_filtered_result:
            display_text_result = format_asr_display_text(text_result) or refined_text_result
            updated_stage_text = merge_transcript_fragments(
                active_segment.get('stage_display_text', ''),
                display_text_result,
            )
            combined_display_text = build_segment_display_text(active_segment, updated_stage_text)
            stable_text, medium_text, partial_text = build_display_layers(
                active_segment,
                "segment_partial",
                stage_text=updated_stage_text,
            )
            if not should_emit_realtime_update(active_segment, combined_display_text, "segment_partial"):
                logger.info(
                    "Skipping noop partial result(segment=%s, chunks=%s, text=%s)",
                    active_segment['segment_id'],
                    active_segment['chunk_count'],
                    refined_text_result[:50],
                )
            else:
                partial_payload = build_realtime_result_payload(
                    session,
                    combined_display_text,
                    chunk_decision,
                    segment_id=active_segment['segment_id'],
                    replace_target_id=active_segment.get('last_result_id'),
                    result_type="segment_partial",
                    stable_text=stable_text,
                    medium_text=medium_text,
                    partial_text=partial_text,
                )
                append_realtime_debug_trace("emit_asr_result", partial_payload)
                emit_realtime_result(sid, partial_payload)
                active_segment['last_result_id'] = partial_payload["result_id"]
                active_segment['latest_display_text'] = combined_display_text
                active_segment['latest_result_type'] = "segment_partial"
                active_segment['stage_display_text'] = updated_stage_text
                logger.info(
                    "ASR partial emitted(result_id=%s, raw=%s, refined=%s, segment=%s, chunks=%s)",
                    partial_payload["result_id"],
                    text_result[:50] if text_result else "",
                    combined_display_text[:50],
                    active_segment['segment_id'],
                    active_segment['chunk_count'],
                )
        elif text_result:
            logger.info(
                "ASR result filtered(raw=%s, refined=%s)",
                text_result[:50],
                refined_text_result[:50] if refined_text_result else "",
            )
            append_realtime_debug_trace(
                "filtered_result",
                {
                    "result_type": "filtered_result",
                    "segment_id": active_segment['segment_id'],
                    "processing_reason": chunk_decision.reason,
                    "text": refined_text_result or text_result,
                },
            )
            if not active_segment.get('last_result_id'):
                logger.info(
                    "Discarding leading filtered realtime segment sid=%s segment=%s chunks=%s duration=%.2fs raw=%s refined=%s",
                    sid,
                    active_segment['segment_id'],
                    active_segment['chunk_count'],
                    active_segment['duration_seconds'],
                    text_result[:50],
                    refined_text_result[:50] if refined_text_result else "",
                )
                append_realtime_debug_trace(
                    "discard_leading_filtered_segment",
                    {
                        "result_type": "discard_leading_filtered_segment",
                        "segment_id": active_segment['segment_id'],
                        "processing_reason": chunk_decision.reason,
                        "text": refined_text_result or text_result,
                    },
                )
                session['active_segment'] = None
                return

        emit_tiered_rewrite_if_needed(
            session,
            sid,
            active_segment,
            chunk_decision,
            finalize_segment=finalize_after_processing,
        )
    except Exception as e:
        logger.exception(f"ASR API call failed: {e}")
        error_payload = {"message": "?????????????"}
        append_realtime_debug_trace("emit_asr_error", error_payload)
        emit("asr_error", error_payload)
    finally:
        session['processing'] = False


def flush_pending_realtime_buffer(session, sid, *, force_finalize_segment=False):
    pending_audio = bytes(session['buffer'])
    if not pending_audio:
        return False

    chunk_decision = DECIDE_STOP_FLUSH(pending_audio, REALTIME_CHUNK_POLICY)
    session['buffer'].clear()

    if not chunk_decision.should_process:
        if chunk_decision.drop_buffer:
            features = chunk_decision.audio_features
            logger.info(
                "Dropping stop flush buffer sid=%s reason=%s speech_gate=%s tail_gate=%s duration=%.2fs rms=%.4f peak=%s active=%.2f voiced=%.2f density=%.2f active_s=%.2fs voiced_s=%.2fs active_run_s=%.2fs voiced_run_s=%.2fs silence=%.2f",
                sid,
                chunk_decision.reason,
                chunk_decision.speech_gate_reason or "-",
                chunk_decision.tail_gate_reason or "-",
                chunk_decision.audio_duration_seconds,
                features.rms if features else 0.0,
                features.peak if features else 0,
                features.active_ratio if features else 0.0,
                features.voiced_ratio if features else 0.0,
                features.voiced_density if features else 0.0,
                features.active_seconds if features else 0.0,
                features.voiced_seconds if features else 0.0,
                features.max_active_run_seconds if features else 0.0,
                features.max_voiced_run_seconds if features else 0.0,
                features.silence_ratio if features else 0.0,
            )
        else:
            retained_audio = retain_realtime_buffer(pending_audio, chunk_decision, REALTIME_CHUNK_POLICY)
            if retained_audio:
                session['buffer'].extend(retained_audio)
                logger.info(
                    "Retaining stop flush buffer sid=%s reason=%s retained_duration=%.2fs original_duration=%.2fs",
                    sid,
                    chunk_decision.reason,
                    len(retained_audio) / float(REALTIME_CHUNK_POLICY.sample_rate * REALTIME_CHUNK_POLICY.bytes_per_sample),
                    len(pending_audio) / float(REALTIME_CHUNK_POLICY.sample_rate * REALTIME_CHUNK_POLICY.bytes_per_sample),
                )
        return False

    process_realtime_audio_chunk(
        session,
        sid,
        pending_audio,
        chunk_decision,
        finalize_after_processing=force_finalize_segment,
    )
    return True


def finalize_active_segment(
    session,
    sid,
    *,
    reason: str,
):
    active_segment = session.get('active_segment')
    if not active_segment:
        return False

    finalize_chunk_decision = ChunkDecision(
        should_process=False,
        reason=reason,
        audio_duration_seconds=active_segment['duration_seconds'],
        trailing_silence_detected=False,
    )
    emitted = emit_tiered_rewrite_if_needed(
        session,
        sid,
        active_segment,
        finalize_chunk_decision,
        finalize_segment=True,
    )
    logger.info(
        "Finalizing realtime segment sid=%s segment=%s reason=%s chunks=%s duration=%.2fs emitted=%s",
        sid,
        active_segment['segment_id'],
        reason,
        active_segment['chunk_count'],
        active_segment['duration_seconds'],
        emitted,
    )
    session['active_segment'] = None
    return emitted


def finalize_active_segment_on_stop(session, sid):
    return finalize_active_segment(
        session,
        sid,
        reason="stop_recording_finalize_segment",
    )


def drain_ready_realtime_buffer(session, sid, *, max_rounds=6):
    rounds = 0
    while rounds < max_rounds:
        if session.get('processing'):
            return

        buffered_audio = bytes(session['buffer'])
        if not buffered_audio:
            break

        chunk_decision = DECIDE_REALTIME_CHUNK(buffered_audio, REALTIME_CHUNK_POLICY)
        if not chunk_decision.should_process:
            if chunk_decision.drop_buffer:
                features = chunk_decision.audio_features
                logger.info(
                    "Dropping realtime buffer sid=%s reason=%s speech_gate=%s tail_gate=%s duration=%.2fs rms=%.4f peak=%s active=%.2f voiced=%.2f density=%.2f active_s=%.2fs voiced_s=%.2fs active_run_s=%.2fs voiced_run_s=%.2fs silence=%.2f",
                    sid,
                    chunk_decision.reason,
                    chunk_decision.speech_gate_reason or "-",
                    chunk_decision.tail_gate_reason or "-",
                    chunk_decision.audio_duration_seconds,
                    features.rms if features else 0.0,
                    features.peak if features else 0,
                    features.active_ratio if features else 0.0,
                    features.voiced_ratio if features else 0.0,
                    features.voiced_density if features else 0.0,
                    features.active_seconds if features else 0.0,
                    features.voiced_seconds if features else 0.0,
                    features.max_active_run_seconds if features else 0.0,
                    features.max_voiced_run_seconds if features else 0.0,
                    features.silence_ratio if features else 0.0,
                )
                session['buffer'].clear()
            else:
                retained_audio = retain_realtime_buffer(buffered_audio, chunk_decision, REALTIME_CHUNK_POLICY)
                if retained_audio != buffered_audio:
                    logger.info(
                        "Retaining realtime buffer sid=%s reason=%s retained_duration=%.2fs original_duration=%.2fs",
                        sid,
                        chunk_decision.reason,
                        len(retained_audio) / float(REALTIME_CHUNK_POLICY.sample_rate * REALTIME_CHUNK_POLICY.bytes_per_sample),
                        len(buffered_audio) / float(REALTIME_CHUNK_POLICY.sample_rate * REALTIME_CHUNK_POLICY.bytes_per_sample),
                    )
                    session['buffer'].clear()
                    session['buffer'].extend(retained_audio)
            break

        session['buffer'].clear()
        process_realtime_audio_chunk(session, sid, buffered_audio, chunk_decision)
        rounds += 1

    if session.get('stop_requested') and not session.get('processing'):
        flush_pending_realtime_buffer(session, sid, force_finalize_segment=True)
        finalize_active_segment_on_stop(session, sid)
        cleanup_realtime_session(sid)


def try_drain_realtime_buffer(session, sid, *, max_rounds=6):
    drain_lock = session.get('drain_lock')
    if drain_lock is None:
        drain_ready_realtime_buffer(session, sid, max_rounds=max_rounds)
        return

    acquired = drain_lock.acquire(blocking=False)
    if not acquired:
        return

    try:
        drain_ready_realtime_buffer(session, sid, max_rounds=max_rounds)
    finally:
        drain_lock.release()


def process_idle_realtime_session(session, sid, *, now: float | None = None) -> bool:
    current_time = time.time() if now is None else now
    last_audio_time = session.get('last_audio_time', session.get('last_process_time', current_time))
    last_speech_time = session.get('last_speech_time', 0.0) or last_audio_time
    last_idle_rewrite_audio_time = session.get('last_idle_rewrite_audio_time', 0.0)

    if session.get('stop_requested') or session.get('processing'):
        return False
    if last_speech_time <= 0 or last_speech_time <= last_idle_rewrite_audio_time:
        return False
    speech_idle_seconds = current_time - last_speech_time
    min_idle_gate_seconds = IDLE_HIGH_REWRITE_SECONDS
    if IDLE_SEGMENT_SPLIT_SECONDS > 0:
        min_idle_gate_seconds = min(IDLE_SEGMENT_SPLIT_SECONDS, IDLE_HIGH_REWRITE_SECONDS)
    if speech_idle_seconds < min_idle_gate_seconds:
        return False

    if IDLE_SEGMENT_SPLIT_SECONDS > 0 and speech_idle_seconds >= IDLE_SEGMENT_SPLIT_SECONDS:
        flush_pending_realtime_buffer(session, sid, force_finalize_segment=False)
        active_segment = session.get('active_segment')
        if not has_displayable_segment_content(active_segment):
            session['last_idle_rewrite_audio_time'] = last_speech_time
            return False

        emitted = finalize_active_segment(
            session,
            sid,
            reason="idle_segment_boundary_timeout",
        )
        session['last_idle_rewrite_audio_time'] = last_speech_time
        logger.info(
            "Idle segment boundary triggered sid=%s speech_idle_seconds=%.2f raw_audio_idle_seconds=%.2f emitted=%s",
            sid,
            speech_idle_seconds,
            current_time - last_audio_time,
            emitted,
        )
        return emitted

    if speech_idle_seconds < IDLE_HIGH_REWRITE_SECONDS:
        return False

    flush_pending_realtime_buffer(session, sid, force_finalize_segment=False)
    active_segment = session.get('active_segment')
    if not active_segment or not active_segment.get('stage_display_text'):
        session['last_idle_rewrite_audio_time'] = last_speech_time
        return False

    idle_chunk_decision = ChunkDecision(
        should_process=False,
        reason="idle_high_rewrite_timeout",
        audio_duration_seconds=active_segment.get('stage_duration_seconds', 0.0),
        trailing_silence_detected=True,
    )
    emitted = emit_stage_rewrite(
        session,
        sid,
        active_segment,
        idle_chunk_decision,
        result_type="high_rewrite",
        processing_reason="idle_high_rewrite_timeout",
        finalize_stage=True,
        ensure_sentence_end=True,
    )
    session['last_idle_rewrite_audio_time'] = last_speech_time
    logger.info(
        "Idle high rewrite triggered sid=%s segment=%s speech_idle_seconds=%.2f raw_audio_idle_seconds=%.2f emitted=%s",
        sid,
        active_segment['segment_id'],
        current_time - last_speech_time,
        current_time - last_audio_time,
        emitted,
    )
    return emitted


def realtime_idle_monitor_loop():
    while True:
        try:
            now = time.time()
            for sid, session in session_mgr.items_snapshot():
                drain_lock = session.get('drain_lock')
                if drain_lock is not None and not drain_lock.acquire(blocking=False):
                    continue
                try:
                    process_idle_realtime_session(session, sid, now=now)
                finally:
                    if drain_lock is not None:
                        drain_lock.release()
        except Exception as exc:
            logger.exception("Realtime idle monitor loop failed: %s", exc)
        time.sleep(1.0)


REALTIME_IDLE_MONITOR_THREAD = threading.Thread(
    target=realtime_idle_monitor_loop,
    name="realtime-idle-monitor",
    daemon=True,
)
REALTIME_IDLE_MONITOR_THREAD.start()


@socketio.on("connect")
def on_connect():
    logger.info(f"Socket.IO 客户端已连接: {request.sid}")


@socketio.on("disconnect")
def on_disconnect():
    sid = request.sid
    logger.info(f"Socket.IO 客户端已断开: {sid}")
    session_mgr.mark_recently_stopped(sid)
    cleanup_realtime_session(sid)


@socketio.on("start_recording")
def on_start_recording(data=None):
    sid = request.sid
    session_mgr.clear_recently_stopped(sid)
    stale_session = session_mgr.get(sid)
    if stale_session and not stale_session.get('processing', False):
        logger.info("收到 start_recording，清理残留 session sid=%s", sid)
        cleanup_realtime_session(sid)
    logger.info("收到 start_recording sid=%s data=%s", sid, data or {})


@socketio.on("audio_stream")
def on_audio_stream(data):
    sid = request.sid
    existing_session = session_mgr.get(sid)
    if existing_session and existing_session.get('stop_requested'):
        logger.info("Ignoring audio_stream after stop request sid=%s bytes=%s", sid, len(data) if isinstance(data, bytes) else "invalid")
        return

    if session_mgr.should_ignore_audio(sid):
        logger.info("Ignoring late audio_stream during stop cooldown sid=%s bytes=%s", sid, len(data) if isinstance(data, bytes) else "invalid")
        return

    session = session_mgr.get_or_create(sid)

    if not isinstance(data, bytes):
        format_error_payload = {"speaker_id": "System", "text": "[数据格式错误]", "is_final": False}
        append_realtime_debug_trace("emit_asr_result", format_error_payload)
        emit("asr_result", format_error_payload)
        return

    fh = session['file_handle']
    fh.write(data)
    fh.flush()

    now = time.time()
    session['last_audio_time'] = now
    if contains_meaningful_realtime_activity(data):
        session['last_speech_time'] = now
    session['buffer'].extend(data)

    try_drain_realtime_buffer(session, sid)


@socketio.on("stop_recording")
def on_stop_recording(data=None):
    sid = request.sid
    session = session_mgr.get(sid)
    session_mgr.mark_recently_stopped(sid)
    if session:
        pending_seconds = len(session['buffer']) / (16000 * 2)
        logger.info(
            "收到 stop_recording sid=%s pending_duration=%.2fs processing=%s",
            sid,
            pending_seconds,
            session.get('processing', False),
        )
        session['stop_requested'] = True

        if session.get('processing', False):
            logger.info("stop_recording 延迟清理 sid=%s，等待当前 chunk 处理完成", sid)
            return

        flush_pending_realtime_buffer(session, sid, force_finalize_segment=True)
        finalize_active_segment_on_stop(session, sid)

    cleanup_realtime_session(sid)


@app.route("/api/v1/llm/summarize", methods=["POST"])
def summarize_meeting():
    data = request.get_json()
    transcript = data.get("transcript", [])
    doc_type = data.get("doc_type", "meeting")

    if not transcript:
        return jsonify({"error": "Transcript is empty"}), 400

    logger.info(f"收到文档生成请求，类型: {doc_type}, 文本条数: {len(transcript)}")

    full_text = build_transcript_text(transcript)
    prompt_config = DOC_PROMPTS.get(doc_type, DOC_PROMPTS["meeting"])

    try:
        response = client.chat.completions.create(
            model=config.LLM_MODEL,
            messages=[
                {"role": "system", "content": prompt_config["system"]},
                {"role": "user", "content": prompt_config["user"].format(full_text=full_text)}
            ],
            stream=False
        )

        raw_content = response.choices[0].message.content
        doc_title, doc_content = parse_doc_response(raw_content, doc_type)

        return jsonify({
            "doc_title": doc_title,
            "summary_text": doc_content,
            "doc_type": doc_type,
            "timestamp": datetime.now().isoformat(),
            "created_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
    except Exception as e:
        logger.error(f"LLM 调用失败: {e}")
        return jsonify({
            "doc_title": "生成失败",
            "summary_text": f"大模型服务连接失败: {str(e)}",
            "error": str(e)
        })


@app.route("/api/v1/llm/regenerate", methods=["POST"])
def regenerate_document():
    data = request.get_json()
    transcript = data.get("transcript", [])
    doc_type = data.get("doc_type", "meeting")
    feedback = data.get("feedback", "")

    if not transcript:
        return jsonify({"error": "Transcript is empty"}), 400
    if not feedback:
        return jsonify({"error": "Feedback is required"}), 400

    logger.info(f"收到文档重新生成请求，类型: {doc_type}, 反馈: {feedback[:50]}...")

    full_text = build_transcript_text(transcript)

    base_prompts = {
        "meeting": {
            "system": "你是一个专业的会议速记员，负责根据转写内容生成会议纪要。",
            "base_user": "请根据以下会议转写内容生成会议纪要。"
        },
        "report": {
            "system": "你是一个专业的商务文案撰写专家，负责根据转写内容生成出差报告。",
            "base_user": "请根据以下出差相关的转写内容生成出差报告。"
        },
        "publicity": {
            "system": "你是一个专业的宣传文案专家，负责根据转写内容生成宣传报道。",
            "base_user": "请根据以下转写内容生成宣传报道。"
        }
    }

    prompt_config = base_prompts.get(doc_type, base_prompts["meeting"])

    enhanced_user_prompt = f"""{prompt_config['base_user']}

【用户反馈要求】
{feedback}

【重要约束】
1. 严格基于输入的转写内容，**禁止编造转写中完全没有提及的事件或成果**
2. 可以对已有信息进行合理扩充：补充背景说明、用更专业/生动的语言重述
3. 如果某类信息在转写中确实未提及，可合理推断或省略

请按照以下格式输出：
【标题】请生成一个简洁准确的标题（10-20字）
【正文】
...文档内容...

转写内容如下：
{full_text}"""

    try:
        response = client.chat.completions.create(
            model=config.LLM_MODEL,
            messages=[
                {"role": "system", "content": prompt_config["system"]},
                {"role": "user", "content": enhanced_user_prompt}
            ],
            stream=False
        )

        raw_content = response.choices[0].message.content
        doc_title, doc_content = parse_doc_response(raw_content, doc_type)

        return jsonify({
            "doc_title": doc_title,
            "summary_text": doc_content,
            "doc_type": doc_type,
            "timestamp": datetime.now().isoformat(),
            "created_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
    except Exception as e:
        logger.error(f"LLM 重新生成失败: {e}")
        return jsonify({
            "doc_title": "重新生成失败",
            "summary_text": f"重新生成失败: {str(e)}",
            "error": str(e)
        })


@app.route("/api/v1/export/word", methods=["POST"])
def export_word():
    data = request.get_json()
    summary = data.get("summary", "")
    transcript = data.get("transcript", [])
    doc_type = data.get("doc_type", "meeting")

    summary = remove_markdown_formatting(summary)

    doc_type_names = {"meeting": "会议纪要", "report": "出差报告", "publicity": "宣传报道"}
    doc_type_name = doc_type_names.get(doc_type, "会议纪要")
    filename = f"{doc_type_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(EXPORT_DIR, filename)

    try:
        doc = Document()

        if doc_type == "meeting":
            doc.add_heading('会议纪要与转写记录', 0)
            doc.add_heading('一、会议总结', level=1)
            doc.add_paragraph(summary)
            doc.add_heading('二、详细转写记录', level=1)
        elif doc_type == "report":
            doc.add_heading('出差报告', 0)
            doc.add_heading('一、报告总结', level=1)
            doc.add_paragraph(summary)
            doc.add_heading('二、详细记录', level=1)
        elif doc_type == "publicity":
            doc.add_heading('宣传报道', 0)
            doc.add_heading('一、宣传内容', level=1)
            doc.add_paragraph(summary)
            doc.add_heading('二、原始记录', level=1)
        else:
            doc.add_heading(f'{doc_type_name}与转写记录', 0)
            doc.add_heading('一、总结', level=1)
            doc.add_paragraph(summary)
            doc.add_heading('二、详细转写记录', level=1)

        for entry in transcript:
            p = doc.add_paragraph()
            speaker = entry.get('speaker', '未知')
            t = entry.get('time', '')
            time_str = f" [{t}]" if t else ""
            p.add_run(f"{speaker}{time_str}: ").bold = True
            p.add_run(entry.get('content', ''))

        doc.save(file_path)
        return jsonify({"download_url": f"/api/v1/export/download/{filename}"})
    except Exception as e:
        logger.error(f"Word 生成失败: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/v1/export/download/<filename>")
def download_file(filename):
    file_path = os.path.join(EXPORT_DIR, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True, download_name=filename)
    return jsonify({"error": "File not found"}), 404


@app.route("/api/v1/audio/upload", methods=["POST"])
def upload_audio():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    file_path = os.path.join(TEMP_DIR, file.filename)
    try:
        content = file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(content)

        logger.info(f"文件上传成功: {file.filename}，开始进行 ASR 识别...")
        text_result = asr_adapter.transcribe_audio_bytes(
            content,
            filename=file.filename or "upload.wav",
            mime_type=file.mimetype or "audio/wav",
        )
        logger.info(f"文件识别完成: {text_result[:50]}...")

        return jsonify({
            "status": "success",
            "filename": file.filename,
            "transcript": text_result,
            "speaker_id": "Speaker_File"
        })
    except Exception as e:
        logger.error(f"文件识别失败: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    logger.info(f"启动服务器: {config.HOST}:{config.PORT}")
    logger.info(f"运行模式: {'内网' if config.USE_INTRANET else '外网开发'}")
    logger.info(f"ASR 模型: {config.ASR_MODEL}, LLM 模型: {config.LLM_MODEL}")
    logger.info(
        "实时转写管线: %s",
        "simplified" if ENABLE_SIMPLIFIED_REALTIME_PIPELINE else "legacy",
    )
    if REALTIME_CHUNK_POLICY_OVERRIDES:
        logger.info("实时转写门槛环境覆盖: %s", REALTIME_CHUNK_POLICY_OVERRIDES)

    ssl_context = None
    if config.ENABLE_HTTPS:
        cert_path = os.path.abspath(config.SSL_CERT)
        key_path = os.path.abspath(config.SSL_KEY)
        if os.path.exists(cert_path) and os.path.exists(key_path):
            ssl_context = (cert_path, key_path)
            logger.info(f"HTTPS 已启用，证书: {cert_path}")
        else:
            logger.warning(f"HTTPS 证书未找到 ({cert_path})，回退到 HTTP 模式")
            logger.warning("请运行 python generate_cert.py 生成证书")

    protocol = "https" if ssl_context else "http"
    logger.info(f"访问地址: {protocol}://localhost:{config.PORT}")

    socketio.run(
        app,
        host=config.HOST,
        port=config.PORT,
        debug=False,
        allow_unsafe_werkzeug=True,
        ssl_context=ssl_context
    )
